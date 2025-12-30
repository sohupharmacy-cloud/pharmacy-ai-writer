import streamlit as st
import google.generativeai as genai
from datetime import datetime
from docx import Document  # 引入 Word 處理工具
import io  # 引入記憶體檔案處理工具

# --- 設定頁面 ---
st.set_page_config(page_title="藥師文案神器 (Word版)", page_icon="💊")
st.title("💊 藥師文案神器 (Word下載版)")
st.caption("輸入主題，AI 幫你寫好並存成 Word 檔！")

# --- 讀取 API Key ---
if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    api_key = st.sidebar.text_input("請輸入 Gemini API Key", type="password")

# --- 輸入區 ---
col1, col2 = st.columns(2)
with col1:
    topic = st.text_input("💡 主題", placeholder="例如：春節腸胃保養")
with col2:
    product = st.text_input("🛍️ 產品/成分", placeholder="例如：益生菌、酵素")

content_points = st.text_area("📝 重點內容", height=150,
                              placeholder="例如：\n1. 肚子脹氣\n2. 飯後散步\n3. 補充好菌")

tone = st.select_slider("🎨 語氣", options=["專業嚴肅", "親切像鄰居", "幽默風趣"], value="親切像鄰居")

# --- 生成按鈕 ---
if st.button("✨ 生成 Word 文案"):
    if not api_key:
        st.error("❌ 請輸入 API Key")
    elif not topic:
        st.warning("⚠️ 請輸入主題")
    else:
        # 設定 AI
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.5-flash')
        
        prompt = f"""
        你是一位台灣藥師。請撰寫一篇社群貼文。
        主題：{topic}
        重點：{content_points}
        產品：{product}
        語氣：{tone}
        規則：符合法規、無療效宣稱、親切、結構清晰、多用Emoji。
        """
        
        with st.spinner("AI 藥師正在打字中..."):
            try:
                response = model.generate_content(prompt)
                final_text = response.text
                
                # 1. 顯示在網頁上預覽
                st.markdown("### 🎉 生成結果預覽：")
                st.markdown("---")
                st.markdown(final_text)
                st.markdown("---")
                
                # 2. 製作 Word 檔 (核心修改處)
                doc = Document()
                doc.add_heading(topic, level=0) # 加入標題
                doc.add_paragraph(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")
                doc.add_paragraph("------------------------------------------------")
                doc.add_paragraph(final_text) # 加入內文
                
                # 把 Word 檔存入記憶體 (而不是硬碟)
                bio = io.BytesIO()
                doc.save(bio)
                
                # 3. 下載按鈕
                file_name = f"{datetime.now().strftime('%Y-%m-%d')}_{topic}.docx"
                
                st.download_button(
                    label="📥 下載 Word 檔 (.docx)",
                    data=bio.getvalue(),
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.success("✅ 準備完成！請點擊上方按鈕下載 Word 檔。")
                
            except Exception as e:
                st.error(f"發生錯誤：{e}")
